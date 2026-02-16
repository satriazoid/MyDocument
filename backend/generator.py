import os
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


# =========================
# REAL WORD TABLE OF CONTENTS
# =========================
def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


# UTILITIES
def safe_filename(text):
    return re.sub(r'[\\/*?:"<>|]', "", text)

def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)  # HITAM WAJIB

def set_spacing(paragraph, space=18):  # 1.5 spasi
    paragraph.paragraph_format.line_spacing = Pt(space)


# MAIN GENERATOR
def generate_doc(title, author, institution, chapters):
    doc = Document()

    # MARGIN KAMPUS
    section = doc.sections[0]
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    # JUDUL (HEADING 1 - TENGAH)
    h1 = doc.add_heading(title.upper(), level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(h1.runs[0], size=14, bold=True)

    p_author = doc.add_paragraph(author)
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_author.runs[0])

    p_inst = doc.add_paragraph(institution)
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p_inst.runs[0])

    # DAFTAR ISI
    doc.add_page_break()
    toc_title = doc.add_heading("DAFTAR ISI", level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(toc_title.runs[0], size=14, bold=True)

    add_table_of_contents(doc)

    # BAB LOOP
    for i in range(1, chapters + 1):
        doc.add_page_break()

        # BAB (HEADING 1 - TENGAH)
        bab = doc.add_heading(f"BAB {i}", level=1)
        bab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(bab.runs[0], size=14, bold=True)

        # HEADING 2 (IKUT BAB)
        h2_title = (
            "Pendahuluan" if i == 1 else
            "Pembahasan" if i == 2 else
            "Penutup"
        )

        h2 = doc.add_heading(f"{i}.1 {h2_title}", level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(h2.runs[0], size=12, bold=True)

        # HEADING 3
        h3 = doc.add_heading(f"{i}.1.1 Uraian", level=3)
        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(h3.runs[0], size=12)

        # ISI PARAGRAF
        p = doc.add_paragraph(
            "Isi makalah di-generate otomatis sesuai standar penulisan ilmiah."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_font(p.runs[0])
        set_spacing(p)  # SPASI 1.5

    # SAVE FILE
    filename = safe_filename(title.replace(" ", "_")) + ".docx"
    os.makedirs("output", exist_ok=True)
    path = os.path.join("output", filename)
    doc.save(path)

    return path
